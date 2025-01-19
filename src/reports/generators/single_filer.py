from src.reports.report_base import BaseReport
from datetime import datetime
from src.reports.data_sources.queries import sql_queries as sql
from docx.enum.text import WD_ALIGN_PARAGRAPH


class SingleFiler(BaseReport):
    """
     Args:
         report_name: name of the report
         template_name: name of template file
         queries: list of query names to execute
         query_params: dict mapping query names to their parameters
         other_data: dict of data directly for the report writer
     """

    def __init__(self, reports_dir, report_name, report_template, queries=None, params=None, other_data=None):
        self.queries = queries or sql
        self.reports_dir = reports_dir
        self.other_data = other_data
        super().__init__(report_name, report_template, self.queries, params)

    def gather_data(self):
        self.data = self.execute_query(self.queries_to_run)

    def _extract_filer_summary(self, replacements):
        for key, value in self.other_data.items():
            replacements[f"${{{key}}}"] = value
        summary = self.data["FilerSummary"][0]
        elements = (0, 2, 3, 4, 5, 6, 7, 8)
        names = ("foundation",
                 "ein",
                 "address",
                 "city",
                 "state",
                 "zip",
                 "tax_year",
                 "assets")
        for index, name in zip(elements, names):
            replacements[f"${{{name}}}"] = str(summary[index])

    def fill_report(self, doc):
        emu = 914400    # 1 in. in English Metric Units (EMU)
        replacements = {
            '${report_title}': 'Summary Report ',
            '${generation_date}': datetime.now().strftime('%Y-%m-%d %H:%M'),
            '${preparer}': 'Don',
        }
        self._extract_filer_summary(replacements)
        # Replace placeholders in paragraphs
        for i, para in enumerate(doc.paragraphs):
            content = para.text
            # Only process paragraphs that contain at least one placeholder
            if '${' in content:
                for key, value in replacements.items():
                    if key in content:
                        content = content.replace(key, value)
                para.text = content

            # Find position to insert table (after "[Contacts]")
            elif "[Contacts]" in para.text:
                para.text = ""  # remove commentary on table insertion
                # Add table after this paragraph
                table = doc.add_table(rows=1, cols=6)
                # Insert table after current paragraph
                para._p.addnext(table._tbl)
                widths = [2*emu, 3*emu, emu, int(.4*emu), emu, 2*emu]
                for idx, width in enumerate(widths):
                    table.columns[idx].width = width
                # table.style = 'Table Grid'            # No settable styles - use defaults

                # Add headers in bold
                headers = table.rows[0].cells
                header_texts = ['Name', 'Address', 'City', 'St', 'Zip', 'Title']
                for idx, text in enumerate(header_texts):
                    cell = headers[idx]
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(text)
                    run.bold = True

                # Add data
                for contact in self.data['Contacts']:
                    row = table.add_row().cells
                    row[0].text = contact[0]
                    row[1].text = contact[1]
                    row[2].text = contact[2]
                    row[3].text = contact[3]
                    row[4].text = contact[4]
                    row[5].text = contact[5]

            elif "[Grants]" in para.text:
                para.text = ""  # remove commentary on table insertion
                # Add table after this paragraph
                table = doc.add_table(rows=1, cols=4)
                widths = [int(.5*emu), 2*emu,  emu, 4*emu]
                for idx, width in enumerate(widths):
                    table.columns[idx].width = width
                # Insert table after current paragraph
                para._p.addnext(table._tbl)
                # table.style = 'Table Grid'            # No settable styles - use defaults

                # Add headers
                headers = table.rows[0].cells
                header_texts = ['TaxYear', 'Recipient', 'Amount', 'Purpose']
                for idx, text in enumerate(header_texts):
                    cell = headers[idx]
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(text)
                    run.bold = True

                # Add data
                for grant in self.data['Grants']:
                    row = table.add_row().cells
                    row[0].text = str(grant[0])
                    row[1].text = grant[1]
                    row[2].text = str(int(grant[3]))
                    row[3].text = grant[2]

                else:
                    para.text = content




